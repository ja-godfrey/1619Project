#%%
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import re

# Load the dataframe
df = pd.read_excel('./../data/derived/03_articles.xlsx')

#TODO: one error throws the whole document out. sometimes saves documents without .docx

# Iterate through each row of the dataframe
for index, row in df.iterrows():
    try:
    
        # Create a new Word document
        doc = Document()
        
        # Add the title to the document
        row['a_title'] = re.sub('[^a-zA-Z\s]+', '', row['a_title'])
        doc.add_heading(row['a_title'], level=1)
        
        # Add the author to the document
        doc.add_heading('Authors:', level=2)
        doc.add_paragraph(row['a_authors'])

        # Add the publisher to the document
        doc.add_heading('Publisher:', level=2)
        doc.add_paragraph(row['media'])

        # Add the publication date to the document
        doc.add_heading('Publication date:', level=2)
        doc.add_paragraph(row['date'])

        # Add the link to the document
        doc.add_heading('Link:', level=2)
        doc.add_paragraph(row['link'])
        
        # Add the image to the document
        # doc.add_picture(row['a_image'])
        doc.add_heading('Splash image:', level=2)
        doc.add_picture(f'./../data/raw/img/{row["clean_title"]}.jpg', width=Inches(4))
        # print(f'./../data/raw/img/{row["clean_title"]}.jpg')
        
        # Add the text to the document
        doc.add_heading('Text:', level=2)
        doc.add_paragraph(row['a_text'])
        
        # Save the document with the title as the file name
        doc.save(f'./../data/derived/articles/{row["a_title"]}.docx')
        print(f'saved #{index}: {row["a_title"]} succesfully')
    except Exception as e:
        print(f'Error {e} on {row["a_title"]} at#{index}')

#%%
# clean the folder
folder_path = './../data/derived/articles'

# get a list of all files in the directory
all_files = os.listdir(folder_path)

dir_path = "./../data/derived/articles/"
for filename in os.listdir(dir_path):
    file_path = os.path.join(dir_path, filename)
    normalized_path = os.path.normpath(file_path)
    if not (filename.endswith(".docx") or filename.endswith(".txt")):
        print(f"Deleting file: {filename}")
        os.remove(normalized_path)

# %%

for index, row in df.iterrows():
    try:
        # Create a new Word document
        doc = Document()

        # Add the title to the document
        row['a_title'] = re.sub('[^a-zA-Z\s]+', '', row['a_title'])
        doc.add_heading(row['a_title'], level=1)

        # Add the author to the document
        doc.add_heading('Authors:', level=2)
        try:
            doc.add_paragraph(row['a_authors'])
        except Exception as e:
            doc.add_paragraph(f'Error retrieving authors: {str(e)}')

        # Add the publisher to the document
        doc.add_heading('Publisher:', level=2)
        try:
            doc.add_paragraph(row['media'])
        except Exception as e:
            doc.add_paragraph(f'Error retrieving publisher: {str(e)}')

        # Add the publication date to the document
        doc.add_heading('Publication date:', level=2)
        try:
            doc.add_paragraph(row['date'])
        except Exception as e:
            doc.add_paragraph(f'Error retrieving publication date: {str(e)}')

        # Add the link to the document
        doc.add_heading('Link:', level=2)
        try:
            doc.add_paragraph(row['link'])
        except Exception as e:
            doc.add_paragraph(f'Error retrieving link: {str(e)}')

        # Add the image to the document
        doc.add_heading('Splash image:', level=2)
        try:
            doc.add_picture(f'./../data/raw/img/{row["clean_title"]}.jpg', width=Inches(4))
        except Exception as e:
            doc.add_paragraph(f'Error retrieving image: {str(e)}')

        # Add the text to the document
        doc.add_heading('Text:', level=2)
        try:
            doc.add_paragraph(row['a_text'])
        except Exception as e:
            doc.add_paragraph(f'Error retrieving text: {str(e)}')

        # Save the document with the title as the file name
        doc.save(f'./../data/derived/articles/{row["a_title"]}.docx')
        print(f'saved #{index}: {row["a_title"]} successfully')
    except Exception as e:
        print(f'Error {e} on {row["a_title"]} at#{index}')


# %%
